"""
Module export transcript ra các định dạng khác nhau
"""
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from datetime import datetime
import io
import streamlit as st

def export_txt(transcript: str, filename: str = "transcript.txt"):
    """Export transcript ra file TXT"""
    return transcript.encode('utf-8'), filename

def export_docx(transcript: str, metadata: dict = None, filename: str = "transcript.docx"):
    """Export transcript ra file DOCX"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Bản Ghi Âm Thanh', 0)
    
    # Metadata
    if metadata:
        doc.add_paragraph(f"Thời gian: {metadata.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}")
        if 'duration' in metadata:
            doc.add_paragraph(f"Độ dài: {format_duration(metadata['duration'])}")
        if 'word_count' in metadata:
            doc.add_paragraph(f"Số từ: {metadata['word_count']}")
        doc.add_paragraph("")  # Empty line
    
    # Transcript content
    paragraphs = transcript.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue(), filename

def export_pdf(transcript: str, metadata: dict = None, filename: str = "transcript.pdf"):
    """Export transcript ra file PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='#1f4e79',
        spaceAfter=12,
        alignment=TA_LEFT
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6
    )
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph("Bản Ghi Âm Thanh", title_style))
    story.append(Spacer(1, 12))
    
    # Metadata
    if metadata:
        meta_text = f"<b>Thời gian:</b> {metadata.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}<br/>"
        if 'duration' in metadata:
            meta_text += f"<b>Độ dài:</b> {format_duration(metadata['duration'])}<br/>"
        if 'word_count' in metadata:
            meta_text += f"<b>Số từ:</b> {metadata['word_count']}<br/>"
        story.append(Paragraph(meta_text, normal_style))
        story.append(Spacer(1, 12))
    
    # Transcript content
    paragraphs = transcript.split('\n')
    for para in paragraphs:
        if para.strip():
            # Escape HTML special characters
            para_escaped = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(para_escaped, normal_style))
            story.append(Spacer(1, 6))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer.getvalue(), filename

def format_duration(seconds: float) -> str:
    """Format duration"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours} giờ {minutes} phút {secs} giây"
    elif minutes > 0:
        return f"{minutes} phút {secs} giây"
    else:
        return f"{secs} giây"

